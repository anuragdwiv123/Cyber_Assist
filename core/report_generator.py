from reportlab.pdfgen import canvas
from docx import Document
import os

def generate_pdf(report_data, filename="cybercrime_report.pdf"):
    c = canvas.Canvas(filename)
    c.drawString(100, 800, "Cybercrime Report")
    y = 760
    for k, v in report_data.items():
        c.drawString(100, y, f"{k}: {v}")
        y -= 20
    c.save()
    return filename

def generate_docx(report_data, filename="cybercrime_report.docx"):
    doc = Document()
    doc.add_heading("Cybercrime Report", 0)
    for k, v in report_data.items():
        doc.add_paragraph(f"{k}: {v}")
    doc.save(filename)
    return filename
